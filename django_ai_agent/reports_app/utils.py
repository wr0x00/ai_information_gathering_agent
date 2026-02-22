from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
from datetime import datetime
from django.conf import settings

def generate_word_report(title, data, template=None):
    """
    Generate a Word document report from data
    
    Args:
        title (str): Report title
        data (dict): Data to include in the report
        template (str, optional): Template content to use
        
    Returns:
        str: Path to the generated report file
    """
    # Create a new Document
    doc = Document()
    
    # Add title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation date
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()  # Empty paragraph for spacing
    
    # If we have a template, use it
    if template:
        doc.add_paragraph(template)
        doc.add_page_break()
    
    # Add data sections
    if isinstance(data, dict):
        for section, content in data.items():
            # Add section heading
            doc.add_heading(section.replace('_', ' ').title(), level=1)
            
            # Add content based on type
            if isinstance(content, str):
                doc.add_paragraph(content)
            elif isinstance(content, list):
                # Create a bulleted list
                for item in content:
                    if isinstance(item, dict):
                        # For dict items, create a table
                        add_dict_to_table(doc, item)
                    else:
                        doc.add_paragraph(str(item), style='List Bullet')
            elif isinstance(content, dict):
                # For nested dicts, create subsections
                for subsection, subcontent in content.items():
                    doc.add_heading(subsection.replace('_', ' ').title(), level=2)
                    if isinstance(subcontent, list):
                        for item in subcontent:
                            if isinstance(item, dict):
                                add_dict_to_table(doc, item)
                            else:
                                doc.add_paragraph(str(item), style='List Bullet')
                    else:
                        doc.add_paragraph(str(subcontent))
            else:
                doc.add_paragraph(str(content))
            
            doc.add_paragraph()  # Empty paragraph for spacing
    
    # Create reports directory if it doesn't exist
    reports_dir = os.path.join(settings.MEDIA_ROOT, 'reports')
    if not os.path.exists(reports_dir):
        os.makedirs(reports_dir)
    
    # Save the document
    filename = f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(reports_dir, filename)
    doc.save(filepath)
    
    return filepath

def add_dict_to_table(doc, data_dict):
    """
    Add a dictionary as a table to the document
    
    Args:
        doc (Document): The Document object
        data_dict (dict): Dictionary to add as a table
    """
    if not data_dict:
        return
    
    # Create table with 2 columns
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Add data to table
    for key, value in data_dict.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)
    
    doc.add_paragraph()  # Empty paragraph for spacing

def generate_combined_report(scan_data, keyword_data, output_filename=None):
    """
    Generate a combined report from scan and keyword data
    
    Args:
        scan_data (dict): Scan results data
        keyword_data (dict): Keyword analysis data
        output_filename (str, optional): Output filename
        
    Returns:
        str: Path to the generated report file
    """
    # Create a new Document
    doc = Document()
    
    # Add title
    heading = doc.add_heading('Combined Information Gathering Report', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation date
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()  # Empty paragraph for spacing
    
    # Add scan data section
    doc.add_heading('Scan Results', level=1)
    if scan_data:
        for module, results in scan_data.items():
            doc.add_heading(module.replace('_', ' ').title(), level=2)
            if isinstance(results, dict):
                add_dict_to_table(doc, results)
            elif isinstance(results, list):
                for item in results:
                    if isinstance(item, dict):
                        add_dict_to_table(doc, item)
                    else:
                        doc.add_paragraph(str(item), style='List Bullet')
            else:
                doc.add_paragraph(str(results))
            doc.add_paragraph()  # Empty paragraph for spacing
    else:
        doc.add_paragraph("No scan data available.")
    
    # Add keyword data section
    doc.add_heading('Keyword Analysis Results', level=1)
    if keyword_data:
        for keyword, results in keyword_data.items():
            doc.add_heading(f"Keyword: {keyword}", level=2)
            if isinstance(results, dict):
                add_dict_to_table(doc, results)
            elif isinstance(results, list):
                for item in results:
                    if isinstance(item, dict):
                        add_dict_to_table(doc, item)
                    else:
                        doc.add_paragraph(str(item), style='List Bullet')
            else:
                doc.add_paragraph(str(results))
            doc.add_paragraph()  # Empty paragraph for spacing
    else:
        doc.add_paragraph("No keyword analysis data available.")
    
    # Create reports directory if it doesn't exist
    reports_dir = os.path.join(settings.MEDIA_ROOT, 'reports')
    if not os.path.exists(reports_dir):
        os.makedirs(reports_dir)
    
    # Determine filename
    if not output_filename:
        output_filename = f"combined_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    filepath = os.path.join(reports_dir, output_filename)
    doc.save(filepath)
    
    return filepath
